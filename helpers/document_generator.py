# helpers/document_generator.py
"""
Generator dokumen Word (.docx) untuk seluruh alur Pokmil.

- Berita Acara Reviu (a)      : dibangun dari nol dengan python-docx (tanpa file template).
- Dokumen Pemilihan / Dokpil (b): mengisi template .docx resmi (docxtpl) via mapping_dokpil.
- Kertas Kerja Evaluasi (c)   : versi Word (versi Excel ada di helpers/excel_generator.py).
- Penetapan Pemenang (d)      : python-docx.
- Laporan Hasil Tender (e)    : python-docx.

Semua fungsi mengembalikan tuple (berhas: bool, hasil: str) di mana hasil adalah
nama file bila berhasil, atau pesan kesalahan bila gagal.
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.config import Config


# ---------- util kecil ----------------------------------------------------

def _rp(nilai):
    try:
        return f"Rp {float(nilai or 0):,.0f}".replace(",", ".")
    except (ValueError, TypeError):
        return "Rp 0"


def _tgl(d):
    bulan = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni",
             "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    if not d:
        return "............................"
    return f"{d.day} {bulan[d.month]} {d.year}"


def _output_path(filename):
    os.makedirs(Config.UPLOAD_FOLDER, exist_ok=True)
    return os.path.join(Config.UPLOAD_FOLDER, filename)


def _safe(s):
    return str(s or "").replace("/", "_").replace("\\", "_").replace(" ", "_")


def _judul(doc, teks):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(teks)
    r.bold = True
    r.font.size = Pt(13)
    return p


def _info_table(doc, pasangan):
    """Tabel 2 kolom (label : nilai) tanpa garis, untuk identitas paket."""
    t = doc.add_table(rows=0, cols=3)
    for label, nilai in pasangan:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = ":"
        row[2].text = str(nilai)
        row[0].width = Cm(5)
        row[1].width = Cm(0.5)
    return t


# ---------- (a) Berita Acara Reviu ----------------------------------------

def generate_docx_ba_reviu(paket, reviu, detail_unsur):
    doc = Document()
    _judul(doc, "BERITA ACARA HASIL REVIU DOKUMEN PERSIAPAN PENGADAAN")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nomor: {reviu.nomor_ba_reviu or '............'}").italic = True

    doc.add_paragraph(
        f"Pada hari ini, {_tgl(reviu.tanggal_ba)}, Pokja Pemilihan bersama Pejabat "
        f"Pembuat Komitmen (PPK) telah melaksanakan reviu atas dokumen persiapan "
        f"pengadaan untuk paket sebagai berikut:"
    )

    _info_table(doc, [
        ("Nama Paket", paket.nama_paket),
        ("Kode Paket / RUP", paket.kode_paket or "-"),
        ("Tahun Anggaran", paket.tahun_anggaran),
        ("Unit Kerja", paket.unit_kerja),
        ("PPK", paket.nama_ppk),
        ("Nilai HPS", _rp(paket.nilai_hps)),
        ("Sumber Dana", paket.sumber_dana),
        ("Jenis Pengadaan", paket.jenis_pengadaan),
        ("Metode Pemilihan", paket.metode_pemilihan),
    ])

    doc.add_paragraph()
    doc.add_paragraph("Hasil reviu terhadap unsur-unsur dokumen persiapan:").bold = True

    tabel = doc.add_table(rows=1, cols=4)
    tabel.style = "Table Grid"
    hdr = tabel.rows[0].cells
    for i, h in enumerate(["No", "Unsur yang Direviu", "Kesesuaian", "Catatan"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    # Kelompokkan per kelompok sesuai data master
    kelompok_urut = ["KAK", "Spesifikasi Teknis", "HPS", "Rancangan Kontrak",
                     "Anggaran", "RUP", "Survei Pasar"]
    no = 1
    grup_terlihat = [k for k in kelompok_urut
                     if any(d.kelompok == k for d in detail_unsur)]
    # tambahkan kelompok lain yang mungkin ada tapi tak terdaftar di urutan
    for d in detail_unsur:
        if d.kelompok not in grup_terlihat:
            grup_terlihat.append(d.kelompok)

    for kel in grup_terlihat:
        baris = tabel.add_row().cells
        baris[0].merge(baris[3])
        baris[0].text = f"Unsur: {kel}"
        baris[0].paragraphs[0].runs[0].bold = True
        for d in [x for x in detail_unsur if x.kelompok == kel]:
            c = tabel.add_row().cells
            c[0].text = str(no)
            c[1].text = d.nama_unsur
            c[2].text = d.kesesuaian or "-"
            c[3].text = d.catatan_reviu or "-"
            no += 1

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Kesimpulan: ").bold = True
    p.add_run(reviu.kesimpulan_reviu or
              "Dokumen persiapan dinyatakan layak untuk dilanjutkan ke tahap pemilihan.")

    # Blok tanda tangan
    doc.add_paragraph()
    ttd = doc.add_table(rows=1, cols=2)
    ttd.rows[0].cells[0].text = "PPK,\n\n\n\n( ........................... )"
    ttd.rows[0].cells[1].text = "Pokja Pemilihan,\n\n\n\n( ........................... )"

    filename = f"BA_Reviu_{_safe(paket.kode_paket or paket.id)}.docx"
    doc.save(_output_path(filename))
    return True, filename


# ---------- (b) Dokumen Pemilihan (isi template MDP) ----------------------

def generate_docx_dokpil(paket, isian=None, template_path=None):
    """Mengisi template MDP menjadi Dokpil.
    - `template_path`: bila diberikan, pakai template itu (dari Manajemen Template).
      Bila None, pilih otomatis via mapping_dokpil dari folder templates_dokumen.
    - `isian` = dict nilai field dari halaman Isian Dokpil (mail-merge LDP/LDK)."""
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        return False, ("Pustaka 'docxtpl' belum terpasang. Jalankan: "
                       "pip install docxtpl")

    if template_path:
        if not os.path.exists(template_path):
            return False, "File template terpilih tidak ditemukan di server."
    else:
        from helpers.mapping_dokpil import tentukan_master_dokpil
        rel = tentukan_master_dokpil(paket)
        template_path = os.path.join(Config.TEMPLATE_DOC_FOLDER, rel)
        if not os.path.exists(template_path):
            return False, (f"Template MDP '{rel}' belum tersedia di folder "
                           f"templates_dokumen. Unggah lewat Manajemen Template "
                           f"atau letakkan file MDP resmi terlebih dahulu.")

    from jinja2 import Environment, Undefined

    class _Kosong(Undefined):
        # Token yang belum diisi dirender sebagai string kosong, bukan error
        def __str__(self):
            return ""

    doc = DocxTemplate(template_path)
    konteks = {
        "nama_paket": paket.nama_paket,
        "kode_paket": paket.kode_paket or "-",
        "tahun_anggaran": paket.tahun_anggaran,
        "unit_kerja": paket.unit_kerja,
        "nama_ppk": paket.nama_ppk,
        "nilai_hps": _rp(paket.nilai_hps),
        "sumber_dana": paket.sumber_dana,
        "jenis_pengadaan": paket.jenis_pengadaan,
        "metode_pemilihan": paket.metode_pemilihan,
        "metode_kualifikasi": paket.metode_kualifikasi,
    }
    # Pemetaan kuratif token LDP -> data paket (auto-fill). Token lain yang tak
    # ada datanya dibiarkan kosong oleh _Kosong untuk diisi Pokja secara manual.
    konteks.update({
        "ldp_kode_rup": paket.kode_paket or "",
        "ldp_nama_paket_pengadaan": paket.nama_paket or "",
        "ldp_uraian_singkat_paket_pengadaan": paket.nama_paket or "",
        "ldp_nama_satuan_kerja_perangkat_daerah": paket.unit_kerja or "",
        "ldp_sumber_dana": paket.sumber_dana or "",
        "ldp_tahun_anggaran": paket.tahun_anggaran or "",
        "ldp_nilai_hps": _rp(paket.nilai_hps),
        "ldp_pejabat_pembuat_komitmen_ppk": paket.nama_ppk or "",
        "ldk_kualifikasi_usaha": "",   # diisi Pokja
    })
    # Nilai dari halaman Isian Dokpil (mail-merge) menimpa/menambah konteks
    if isian:
        konteks.update({k: v for k, v in isian.items() if v not in (None, "")})
    doc.render(konteks, jinja_env=Environment(undefined=_Kosong))
    filename = f"Dokpil_{_safe(paket.kode_paket or paket.id)}.docx"
    doc.save(_output_path(filename))
    return True, filename


# ---------- (c) Kertas Kerja Evaluasi (versi Word) ------------------------

def generate_docx_kertas_kerja(paket, peserta_list, kriteria_list=None):
    kriteria_list = kriteria_list or []
    teknis = [k for k in kriteria_list if k.kelompok == 'Teknis']
    kualif = [k for k in kriteria_list if k.kelompok == 'Kualifikasi']

    doc = Document()
    _judul(doc, "KERTAS KERJA EVALUASI")
    doc.add_paragraph(f"Paket: {paket.nama_paket}").bold = True
    doc.add_paragraph(f"Metode Pemilihan: {paket.metode_pemilihan} | HPS: {_rp(paket.nilai_hps)}")

    def _tabel_kriteria(judul, kriteria, dengan_bobot, sumber):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(judul).bold = True
        doc.add_paragraph(f"(disusun dari {sumber})").italic = True
        kolom = (["No", "Kriteria", "Bobot (%)", "Ambang"] if dengan_bobot
                 else ["No", "Persyaratan"]) + [ps.nama_penyedia for ps in peserta_list]
        t = doc.add_table(rows=1, cols=len(kolom)); t.style = "Table Grid"
        for i, h in enumerate(kolom):
            t.rows[0].cells[i].text = h
            if t.rows[0].cells[i].paragraphs[0].runs:
                t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for j, k in enumerate(kriteria, 1):
            c = t.add_row().cells
            c[0].text = str(j)
            c[1].text = k.nama_kriteria
            off = 2
            if dengan_bobot:
                c[2].text = f"{k.bobot or 0:g}"
                c[3].text = f"{k.ambang_batas or 0:g}"
                off = 4
            for i in range(len(peserta_list)):
                c[off + i].text = ""   # diisi Pokja
        if not kriteria:
            doc.add_paragraph("Kriteria belum diisi pada menu Setting Kriteria.").italic = True

    _tabel_kriteria("A. Evaluasi Teknis", teknis, True, "LKE pada LDP")
    _tabel_kriteria("B. Evaluasi Kualifikasi", kualif, False, "persyaratan pada LDK")

    # C. Rekap harga
    doc.add_paragraph()
    doc.add_paragraph("C. Rekapitulasi Harga").bold = True
    th = doc.add_table(rows=1, cols=4); th.style = "Table Grid"
    for i, h in enumerate(["No", "Penyedia", "Penawaran", "Terkoreksi"]):
        th.rows[0].cells[i].text = h
        th.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, ps in enumerate(peserta_list, 1):
        c = th.add_row().cells
        c[0].text = str(i); c[1].text = ps.nama_penyedia
        c[2].text = _rp(ps.harga_penawaran); c[3].text = _rp(ps.harga_terkoreksi)

    filename = f"KertasKerja_{_safe(paket.kode_paket or paket.id)}.docx"
    doc.save(_output_path(filename))
    return True, filename


# ---------- (d) Penetapan Pemenang ----------------------------------------

def generate_docx_penetapan(paket, penetapan):
    doc = Document()
    _judul(doc, "PENETAPAN PEMENANG")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nomor: {penetapan.nomor_sk or '............'}").italic = True

    doc.add_paragraph(
        f"Berdasarkan hasil evaluasi penawaran untuk paket "
        f"\"{paket.nama_paket}\" (HPS {_rp(paket.nilai_hps)}), Pokja Pemilihan "
        f"menetapkan pemenang sebagai berikut:"
    )
    pem = penetapan.pemenang
    cad = penetapan.cadangan
    _info_table(doc, [
        ("Pemenang", pem.nama_penyedia if pem else "-"),
        ("NPWP", (pem.npwp if pem else "-") or "-"),
        ("Harga Terkoreksi", _rp(pem.harga_terkoreksi) if pem else "-"),
        ("Pemenang Cadangan", cad.nama_penyedia if cad else "-"),
        ("Tanggal Penetapan", _tgl(penetapan.tanggal_sk)),
    ])
    if penetapan.dasar_penetapan:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Dasar Penetapan: ").bold = True
        p.add_run(penetapan.dasar_penetapan)

    doc.add_paragraph()
    ttd = doc.add_table(rows=1, cols=1)
    ttd.rows[0].cells[0].text = "Pokja Pemilihan,\n\n\n\n( ........................... )\nKetua"

    filename = f"Penetapan_{_safe(paket.kode_paket or paket.id)}.docx"
    doc.save(_output_path(filename))
    return True, filename


# ---------- (e) Laporan Hasil Tender --------------------------------------

def generate_docx_laporan(paket, laporan, peserta_list, penetapan=None):
    doc = Document()
    _judul(doc, "LAPORAN HASIL TENDER")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Nomor: {laporan.nomor_laporan or '............'} | "
              f"Tanggal: {_tgl(laporan.tanggal)}").italic = True

    doc.add_paragraph("A. Informasi Paket").bold = True
    _info_table(doc, [
        ("Nama Paket", paket.nama_paket),
        ("Kode Paket / RUP", paket.kode_paket or "-"),
        ("Unit Kerja", paket.unit_kerja),
        ("HPS", _rp(paket.nilai_hps)),
        ("Metode Pemilihan", paket.metode_pemilihan),
    ])

    doc.add_paragraph()
    doc.add_paragraph("B. Peserta dan Hasil Evaluasi").bold = True
    tabel = doc.add_table(rows=1, cols=5)
    tabel.style = "Table Grid"
    for i, h in enumerate(["No", "Penyedia", "Terkoreksi", "Hasil", "Catatan"]):
        sel = tabel.rows[0].cells[i]
        sel.text = h
        sel.paragraphs[0].runs[0].bold = True
    for i, ps in enumerate(peserta_list, start=1):
        c = tabel.add_row().cells
        c[0].text = str(i)
        c[1].text = ps.nama_penyedia
        c[2].text = _rp(ps.harga_terkoreksi)
        c[3].text = "LULUS" if ps.lulus_semua else "GUGUR"
        c[4].text = ps.catatan or "-"

    doc.add_paragraph()
    doc.add_paragraph("C. Kesimpulan").bold = True
    pemenang = penetapan.pemenang if penetapan and penetapan.pemenang else None
    doc.add_paragraph(
        laporan.ringkasan or
        (f"Pemenang tender adalah {pemenang.nama_penyedia} dengan harga "
         f"{_rp(pemenang.harga_terkoreksi)}." if pemenang else
         "Hasil tender sebagaimana tabel di atas.")
    )

    filename = f"Laporan_{_safe(paket.kode_paket or paket.id)}.docx"
    doc.save(_output_path(filename))
    return True, filename
