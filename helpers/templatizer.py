# helpers/templatizer.py
"""Mengubah dokumen (.docx) menjadi template mail-merge docxtpl. Memindai titik
isian (garis bawah '____', titik '....', atau instruksi '[diisi ...]') baik di
TABEL maupun PARAGRAF. Default fokus LDP & LDK; bisa 'semua' bagian."""
import re
import shutil
import unicodedata
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

TOKEN = re.compile(r'{{\s*([a-zA-Z_]\w*)\s*}}')
ISIAN = re.compile(
    r'(?:_{3,}|\.{4,})\s*(?:\[[^\]]*?(?:diisi|coret|contoh|pilih|apabila)[^\]]*?\])?'
    r'|\[[^\]]*?diisi[^\]]*?\]', re.I)
HINT = re.compile(r'\[[^\]]*?(?:coret|contoh|pilih|apabila)[^\]]*?\]', re.I)
INSTR_PENUH = re.compile(r'^\[[^\]]*?(?:diisi|coret|contoh|pilih|apabila)[^\]]*?\]$', re.I)


def _slug(teks, maks=34):
    teks = unicodedata.normalize('NFKD', teks).encode('ascii', 'ignore').decode().lower()
    teks = re.sub(r'[^a-z0-9]+', '_', teks).strip('_')
    teks = re.sub(r'^[0-9_]+', '', teks)
    return teks[:maks].strip('_')


def _nama_konteks(pre):
    pre = re.split(r'[.;]\s', pre)[-1]
    kata = re.findall(r'[A-Za-z]+', pre)[-6:]
    return _slug(" ".join(kata)) or "isian"


def _proses_paragraf(p, prefix, dipakai, fields):
    s = p.text.strip()
    if INSTR_PENUH.match(s):
        for r in p.runs:
            r.text = ""
        return
    teks = p.text
    if not ISIAN.search(teks):
        return
    out, last = [], 0
    for m in ISIAN.finditer(teks):
        out.append(teks[last:m.start()])
        dasar = f"{prefix}_{_nama_konteks(teks[:m.start()])}"
        n = dipakai.get(dasar, 0) + 1
        dipakai[dasar] = n
        tok = dasar if n == 1 else f"{dasar}_{n}"
        fields.append(tok)
        out.append("{{ " + tok + " }}")
        last = m.end()
    out.append(teks[last:])
    teks = re.sub(r'[ \t]{2,}', ' ', HINT.sub('', "".join(out))).strip()
    if p.runs:
        p.runs[0].text = teks
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(teks)


def _proses_sel(sel, prefix, dipakai, fields):
    for p in sel.paragraphs:
        _proses_paragraf(p, prefix, dipakai, fields)


def _ekstrak_token_manual(doc):
    """Kumpulkan token {{ }} yang sudah ditandai manual (paragraf & sel tabel)."""
    teks_all = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                teks_all.append(c.text)
    seen, urut = set(), []
    for teks in teks_all:
        for k in TOKEN.findall(teks):
            if k not in seen:
                seen.add(k); urut.append(k)
    return urut


def templatize_docx(src_path, dst_path, lingkup='ldp_ldk'):
    """Tokenisasi LDP & LDK (default) atau 'semua' bagian ber-isian (tabel & paragraf).
    Jika dokumen SUDAH ditandai {{ }} secara manual, token itu langsung dipakai
    (lingkup='manual' memaksa mode ini). Mengembalikan daftar key field unik."""
    doc = Document(src_path)

    # Prioritas: bila sudah ada token {{ }} manual, pakai apa adanya (template siap).
    manual = _ekstrak_token_manual(doc)
    if lingkup == 'manual' or manual:
        shutil.copyfile(src_path, dst_path)   # simpan tanpa mengubah apa pun
        return manual

    fields, dipakai = [], {}
    kode = None
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p = Paragraph(child, doc)
            t = p.text.strip().upper()
            if t.endswith('(LDP)'):
                kode = 'ldp'
            elif t.endswith('(LDK)'):
                kode = 'ldk'
            elif t.startswith('BAB '):
                kode = None
            aktif = kode if lingkup == 'ldp_ldk' else (kode or 'isian')
            if (lingkup == 'semua' or kode in ('ldp', 'ldk')) and aktif:
                _proses_paragraf(p, aktif, dipakai, fields)
        elif tag == 'tbl':
            aktif = kode if lingkup == 'ldp_ldk' else (kode or 'isian')
            if (lingkup == 'semua' or kode in ('ldp', 'ldk')) and aktif:
                tbl = Table(child, doc)
                if any(ISIAN.search(c.text) for r in tbl.rows for c in r.cells):
                    for r in tbl.rows:
                        target = r.cells[1:] if len(r.cells) > 1 else r.cells
                        for c in target:
                            _proses_sel(c, aktif, dipakai, fields)
    doc.save(dst_path)
    seen, urut = set(), []
    for f in fields:
        if f not in seen:
            seen.add(f); urut.append(f)
    return urut


def label_dari_key(key):
    grup = 'LDP' if key.startswith('ldp_') else ('LDK' if key.startswith('ldk_') else 'Isian')
    inti = re.sub(r'^(ldp_|ldk_|isian_)', '', key).replace('_', ' ').strip()
    return grup, (inti[:1].upper() + inti[1:]) if inti else key
