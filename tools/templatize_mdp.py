#!/usr/bin/env python3
"""
templatize_mdp.py — mengubah MDP LKPP (Perlem 12/2021) menjadi template
mail-merge docxtpl, KHUSUS bagian LDP dan LDK (SSKK/Rancangan Kontrak tidak
disentuh karena merupakan dokumen PPK).

Setiap titik isian (garis bawah '____' dan/atau instruksi '[diisi ...]') pada
tabel LDP/LDK diganti token {{ ldp_namafield }} / {{ ldk_namafield }}. Nama
field diturunkan dari teks sebelum titik isian agar mudah dikenali.

Pemakaian:
    python templatize_mdp.py "MDP Asli.docx" "Template.docx"
"""
import re, sys, unicodedata
from docx import Document
from docx.table import Table

# titik isian: garis bawah (+instruksi opsional) ATAU instruksi "diisi" sendiri
ISIAN = re.compile(
    r'(?:_{3,}|\.{4,})\s*(?:\[[^\]]*?(?:diisi|coret|contoh|pilih|apabila)[^\]]*?\])?'
    r'|\[[^\]]*?diisi[^\]]*?\]', re.I)
# instruksi-hint tanpa garis bawah (contoh/apabila/pilih/coret) -> dibuang saja
HINT = re.compile(r'\[[^\]]*?(?:coret|contoh|pilih|apabila)[^\]]*?\]', re.I)
INSTR_PENUH = re.compile(r'^\[[^\]]*?(?:diisi|coret|contoh|pilih|apabila)[^\]]*?\]$', re.I)


def slug(teks, maks=34):
    teks = unicodedata.normalize('NFKD', teks).encode('ascii', 'ignore').decode().lower()
    teks = re.sub(r'[^a-z0-9]+', '_', teks).strip('_')
    teks = re.sub(r'^[0-9_]+', '', teks)
    return teks[:maks].strip('_')


def nama_dari_konteks(pre):
    """Ambil ~6 kata terakhir sebelum titik isian sebagai nama field."""
    pre = re.split(r'[.;]\s', pre)[-1]          # kalimat terakhir
    kata = re.findall(r'[A-Za-z]+', pre)[-6:]
    return slug(" ".join(kata)) or "isian"


def proses_sel(sel, prefix, dipakai, fields):
    dibuat = 0
    for p in sel.paragraphs:
        s = p.text.strip()
        if INSTR_PENUH.match(s):                # paragraf instruksi murni -> hapus
            for r in p.runs:
                r.text = ""
            continue
        teks = p.text
        if not ISIAN.search(teks):
            continue
        out, last = [], 0
        for m in ISIAN.finditer(teks):
            out.append(teks[last:m.start()])
            dasar = f"{prefix}_{nama_dari_konteks(teks[:m.start()])}"
            n = dipakai.get(dasar, 0) + 1
            dipakai[dasar] = n
            tok = dasar if n == 1 else f"{dasar}_{n}"
            fields.append(tok); dibuat += 1
            out.append("{{ " + tok + " }}")
            last = m.end()
        out.append(teks[last:])
        teks = HINT.sub('', "".join(out))
        teks = re.sub(r'[ \t]{2,}', ' ', teks).strip()
        if p.runs:
            p.runs[0].text = teks
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.add_run(teks)
    return dibuat


def templatize(src, dst):
    doc = Document(src)
    fields, dipakai, total = [], {}, 0
    kode = None
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            t = ''.join(n.text or '' for n in child.iter()
                        if n.tag.endswith('}t')).strip().upper()
            if t.endswith('(LDP)'):
                kode = 'ldp'
            elif t.endswith('(LDK)'):
                kode = 'ldk'
            elif t.startswith('BAB '):
                kode = None
        elif tag == 'tbl' and kode in ('ldp', 'ldk'):
            tbl = Table(child, doc)
            for r in tbl.rows:
                target = r.cells[1:] if len(r.cells) > 1 else r.cells
                for c in target:
                    total += proses_sel(c, kode, dipakai, fields)
    doc.save(dst)
    return fields, total


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print(__doc__); sys.exit(1)
    fields, total = templatize(sys.argv[1], sys.argv[2])
    open(sys.argv[2].rsplit('.', 1)[0] + ".fields.txt", "w",
         encoding="utf-8").write("\n".join(sorted(set(fields))))
    ldp = len({f for f in fields if f.startswith('ldp_')})
    ldk = len({f for f in fields if f.startswith('ldk_')})
    print(f"OK {sys.argv[2].split('/')[-1][:55]:55} : LDP={ldp:3d} LDK={ldk:3d} (total {total})")
